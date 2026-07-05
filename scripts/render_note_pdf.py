#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import html
import importlib.util
import json
import re
import sys
import tempfile
from pathlib import Path
from typing import Any


CSS = """
@page {
  size: A4;
  margin: 20mm 17mm;
}
body {
  max-width: 860px;
  margin: 0 auto;
  padding: 32px 22px 56px;
  color: #1f2933;
  background: #ffffff;
  font-family: "Microsoft YaHei", "Noto Sans CJK SC", "Source Han Sans SC", "PingFang SC", "SimSun", Arial, sans-serif;
  line-height: 1.72;
  font-size: 15.5px;
}
h1, h2, h3, h4 {
  color: #111827;
  line-height: 1.35;
  break-after: avoid;
}
h1 {
  font-size: 28px;
  border-bottom: 2px solid #d0d7de;
  padding-bottom: 10px;
}
h2 {
  margin-top: 30px;
  font-size: 21px;
  border-bottom: 1px solid #e5e7eb;
  padding-bottom: 6px;
}
h3 {
  font-size: 17px;
}
p, li {
  text-align: justify;
}
a {
  color: #0f766e;
}
blockquote {
  margin: 16px 0;
  padding: 10px 14px;
  color: #374151;
  background: #f6f8fa;
  border-left: 4px solid #9ca3af;
}
code {
  font-family: Consolas, "Courier New", monospace;
  background: #f6f8fa;
  padding: 0.12em 0.32em;
  border-radius: 4px;
}
pre {
  padding: 12px;
  overflow-x: auto;
  background: #f6f8fa;
  border: 1px solid #e5e7eb;
  border-radius: 6px;
}
pre code {
  background: transparent;
  padding: 0;
}
figure {
  margin: 20px auto;
  break-inside: avoid;
}
img {
  display: block;
  max-width: 100%;
  height: auto;
  margin: 0 auto;
  border: 1px solid #d1d5db;
}
figcaption {
  margin-top: 8px;
  color: #4b5563;
  font-size: 13px;
  text-align: center;
}
table {
  width: 100%;
  border-collapse: collapse;
  margin: 16px 0;
  break-inside: avoid;
}
th, td {
  border: 1px solid #d1d5db;
  padding: 6px 8px;
  vertical-align: top;
}
th {
  background: #f3f4f6;
}
"""


def module_available(name: str) -> bool:
    return importlib.util.find_spec(name) is not None



def convert_with_markdown_module(markdown_text: str) -> str | None:
    if not module_available("markdown"):
        return None
    try:
        import markdown  # type: ignore

        return markdown.markdown(
            markdown_text,
            extensions=["extra", "tables", "fenced_code", "sane_lists", "nl2br"],
            output_format="html5",
        )
    except Exception:
        return None


def inline_markdown(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", lambda m: f"<code>{m.group(1)}</code>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", escaped)
    escaped = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        lambda m: f'<a href="{html.escape(m.group(2), quote=True)}">{m.group(1)}</a>',
        escaped,
    )
    return escaped


def render_table(lines: list[str]) -> str:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if not rows:
        return ""
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[1]) else rows[1:]
    parts = ["<table>", "<thead><tr>"]
    for cell in header:
        parts.append(f"<th>{inline_markdown(cell)}</th>")
    parts.append("</tr></thead>")
    if body:
        parts.append("<tbody>")
        for row in body:
            parts.append("<tr>")
            for cell in row:
                parts.append(f"<td>{inline_markdown(cell)}</td>")
            parts.append("</tr>")
        parts.append("</tbody>")
    parts.append("</table>")
    return "\n".join(parts)


def fallback_markdown_to_html(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    out: list[str] = []
    paragraph: list[str] = []
    list_items: list[str] = []
    table_lines: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            out.append(f"<p>{inline_markdown(' '.join(paragraph))}</p>")
            paragraph = []

    def flush_list() -> None:
        nonlocal list_items
        if list_items:
            out.append("<ul>")
            for item in list_items:
                out.append(f"<li>{inline_markdown(item)}</li>")
            out.append("</ul>")
            list_items = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            out.append(render_table(table_lines))
            table_lines = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.startswith("```"):
            flush_paragraph()
            flush_list()
            flush_table()
            if in_code:
                out.append(f"<pre><code>{html.escape(chr(10).join(code_lines))}</code></pre>")
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph()
            flush_list()
            flush_table()
            continue

        if line.lstrip().startswith("|") and "|" in line.strip()[1:]:
            flush_paragraph()
            flush_list()
            table_lines.append(line)
            continue
        else:
            flush_table()

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if image_match:
            flush_paragraph()
            flush_list()
            alt = html.escape(image_match.group(1))
            src = html.escape(image_match.group(2), quote=True)
            out.append(f'<figure><img src="{src}" alt="{alt}"><figcaption>{alt}</figcaption></figure>')
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            flush_list()
            level = len(heading.group(1))
            out.append(f"<h{level}>{inline_markdown(heading.group(2).strip())}</h{level}>")
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            list_items.append(bullet.group(1).strip())
            continue

        quote = re.match(r"^\s*>\s?(.+)$", line)
        if quote:
            flush_paragraph()
            flush_list()
            out.append(f"<blockquote>{inline_markdown(quote.group(1).strip())}</blockquote>")
            continue

        paragraph.append(line.strip())

    flush_paragraph()
    flush_list()
    flush_table()
    if in_code:
        out.append(f"<pre><code>{html.escape(chr(10).join(code_lines))}</code></pre>")
    return "\n".join(out)


def plain_text(markdown_text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1", markdown_text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("\\|", "|")
    return text.strip()


def markdown_to_html(markdown_text: str) -> tuple[str, str]:
    html_body = convert_with_markdown_module(markdown_text)
    if html_body is not None:
        return html_body, "python-markdown"
    return fallback_markdown_to_html(markdown_text), "fallback"


def html_document(body: str, title: str) -> str:
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>{CSS}</style>
</head>
<body>
{body}
</body>
</html>
"""



def infer_title(markdown_text: str, note_path: Path) -> str:
    for line in markdown_text.splitlines():
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return match.group(1).strip()
    return note_path.stem


DOCX_FONT_NAME = "Microsoft YaHei"


def set_rfonts(element: Any, font_name: str = DOCX_FONT_NAME) -> None:
    try:
        get_or_add_rpr = getattr(element, "get_or_add_rPr", None)
        rpr = get_or_add_rpr() if callable(get_or_add_rpr) else getattr(element, "rPr", None)
        if rpr is None:
            return
        rfonts = getattr(rpr, "rFonts", None)
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")  # type: ignore[name-defined]
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), font_name)  # type: ignore[name-defined]
    except Exception:
        pass


def set_docx_font(run: Any, font_name: str = DOCX_FONT_NAME) -> None:
    try:
        run.font.name = font_name
        set_rfonts(run._element, font_name)
    except Exception:
        pass


def set_document_fonts(document: Any, font_name: str = DOCX_FONT_NAME) -> None:
    for style in document.styles:
        try:
            style.font.name = font_name
            set_rfonts(style._element, font_name)
        except Exception:
            pass


def enforce_document_fonts(document: Any, font_name: str = DOCX_FONT_NAME) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            set_docx_font(run, font_name)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_docx_font(run, font_name)

    for section in document.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                for run in paragraph.runs:
                    set_docx_font(run, font_name)

def add_docx_paragraph(document: Any, text: str, style: str | None = None) -> Any:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    run = paragraph.add_run(plain_text(text))
    set_docx_font(run)
    return paragraph


def add_docx_table(document: Any, lines: list[str]) -> None:
    rows = [[plain_text(cell.strip()) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in rows[1]):
        rows = [rows[0]] + rows[2:]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            run = paragraph.add_run(cell_text)
            run.bold = row_idx == 0
            set_docx_font(run)


def add_docx_image(document: Any, note_dir: Path, alt_text: str, image_ref: str) -> None:
    temp_path: Path | None = None
    try:
        if image_ref.startswith("data:image/"):
            header, encoded = image_ref.split(",", 1)
            mime = header[5:].split(";", 1)[0].lower()
            suffix = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/jpg": ".jpg",
                "image/gif": ".gif",
                "image/webp": ".webp",
            }.get(mime, ".img")
            data = base64.b64decode(encoded)
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(data)
                temp_path = Path(tmp.name)
            image_path = temp_path
        else:
            image_path = (note_dir / image_ref).resolve()
            if not image_path.exists():
                add_docx_paragraph(document, f"[Image placeholder: {alt_text}. Missing file: {image_ref}]")
                return

        document.add_picture(str(image_path), width=Inches(6.2))  # type: ignore[name-defined]
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[name-defined]
        run = caption.add_run(plain_text(alt_text))
        run.italic = True
        set_docx_font(run)
    except Exception as exc:
        add_docx_paragraph(document, f"[Image placeholder: {alt_text}. DOCX insert failed: {exc}]")
    finally:
        if temp_path is not None:
            try:
                temp_path.unlink(missing_ok=True)
            except Exception:
                pass


def render_docx(markdown_text: str, note_path: Path, docx_path: Path) -> tuple[bool, str]:
    if not module_available("docx"):
        return False, "python-docx is not installed."
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml import OxmlElement as _OxmlElement  # type: ignore
        from docx.oxml.ns import qn as _qn  # type: ignore
        from docx.shared import Inches as _Inches  # type: ignore

        globals()["Document"] = Document
        globals()["Inches"] = _Inches
        globals()["WD_ALIGN_PARAGRAPH"] = _WD_ALIGN_PARAGRAPH
        globals()["OxmlElement"] = _OxmlElement
        globals()["qn"] = _qn
    except Exception as exc:
        return False, f"python-docx import failed: {exc}"

    document = Document()
    set_document_fonts(document)
    note_dir = note_path.parent
    lines = markdown_text.splitlines()
    index = 0
    paragraph_parts: list[str] = []
    bullet_items: list[str] = []
    table_lines: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_parts
        if paragraph_parts:
            add_docx_paragraph(document, " ".join(paragraph_parts))
            paragraph_parts = []

    def flush_bullets() -> None:
        nonlocal bullet_items
        for item in bullet_items:
            add_docx_paragraph(document, item, style="List Bullet")
        bullet_items = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_docx_table(document, table_lines)
            table_lines = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_bullets()
            flush_table()
            if in_code:
                add_docx_paragraph(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            flush_bullets()
            flush_table()
            index += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            flush_bullets()
            table_lines.append(stripped)
            index += 1
            continue
        flush_table()

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if image_match:
            flush_paragraph()
            flush_bullets()
            add_docx_image(document, note_dir, image_match.group(1), image_match.group(2))
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            flush_bullets()
            level = min(4, len(heading.group(1)))
            paragraph = document.add_heading(plain_text(heading.group(2)), level=level)
            for run in paragraph.runs:
                set_docx_font(run)
            index += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            bullet_items.append(bullet.group(1).strip())
            index += 1
            continue

        quote = re.match(r"^\s*>\s?(.+)$", stripped)
        if quote:
            flush_paragraph()
            flush_bullets()
            add_docx_paragraph(document, quote.group(1).strip(), style="Intense Quote" if "Intense Quote" in document.styles else None)
            index += 1
            continue

        paragraph_parts.append(stripped)
        index += 1

    flush_paragraph()
    flush_bullets()
    flush_table()
    if in_code and code_lines:
        add_docx_paragraph(document, "\n".join(code_lines))

    try:
        enforce_document_fonts(document)
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(docx_path))
    except Exception as exc:
        return False, f"DOCX save failed: {exc}"
    return docx_path.exists() and docx_path.stat().st_size > 0, "python-docx"


def safe_output_basename(value: str) -> str:
    name = Path(value).name
    if Path(name).suffix.lower() in {".md", ".html", ".docx"}:
        name = Path(name).stem
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name).strip(" ._")
    if not name:
        raise ValueError("Output basename cannot be empty.")
    return name


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a paper note Markdown file to DOCX, with optional HTML debug output.")
    parser.add_argument("note_md", help="Path to the Markdown note.")
    parser.add_argument("--output-dir", help="Directory for DOCX output. Defaults to the Markdown file directory.")
    parser.add_argument("--output-basename", help="Base filename without extension for DOCX output. Defaults to the Markdown filename stem.")
    parser.add_argument("--docx-name", help="DOCX output filename. Overrides --output-basename for DOCX.")
    parser.add_argument("--write-html", action="store_true", help="Also write an HTML debug copy. Default final outputs are Markdown and DOCX only.")
    parser.add_argument("--html-name", help="HTML output filename when --write-html is used.")
    parser.add_argument("--title", help="HTML title when --write-html is used. Defaults to first H1 or note filename.")
    parser.add_argument("--skip-docx", action="store_true", help="Do not write DOCX.")
    parser.add_argument("--write-report", action="store_true", help="Write render_report.json for debugging.")
    args = parser.parse_args()

    note_path = Path(args.note_md).expanduser().resolve()
    if not note_path.exists():
        print(f"Markdown note not found: {note_path}", file=sys.stderr)
        return 1

    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else note_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    try:
        output_basename = safe_output_basename(args.output_basename or note_path.stem)
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1

    docx_path = output_dir / (args.docx_name or f"{output_basename}.docx")
    html_path = output_dir / (args.html_name or f"{output_basename}.html")

    markdown_text = note_path.read_text(encoding="utf-8")
    report: dict[str, Any] = {
        "note_md": str(note_path),
        "note_docx": str(docx_path),
        "output_basename": output_basename,
        "docx_created": False,
        "docx_backend": None,
        "html_created": False,
        "note_html": None,
        "markdown_backend": None,
        "warnings": [],
    }

    if args.write_html:
        body, markdown_backend = markdown_to_html(markdown_text)
        title = args.title or infer_title(markdown_text, note_path)
        html_path.write_text(html_document(body, title), encoding="utf-8")
        report["html_created"] = True
        report["note_html"] = str(html_path)
        report["markdown_backend"] = markdown_backend

    if args.skip_docx:
        report["warnings"].append("DOCX rendering skipped by --skip-docx.")
    else:
        ok, backend = render_docx(markdown_text, note_path, docx_path)
        report["docx_created"] = ok
        report["docx_backend"] = backend if ok else None
        if not ok:
            report["warnings"].append(f"DOCX was not generated: {backend}")

    if args.write_report:
        (output_dir / "render_report.json").write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")

    if report["html_created"]:
        print(f"HTML debug copy written: {html_path}")
    if report["docx_created"]:
        print(f"DOCX written: {docx_path}")
    else:
        print("DOCX was not generated.")

    for warning in report["warnings"]:
        print(f"- {warning}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())

